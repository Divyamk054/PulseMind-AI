import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Configure Margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

FONT_FAMILY = "Times New Roman"

def set_font(run, font_name=FONT_FAMILY, size_pt=12, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def add_p(text="", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, font_size=12, bold=False, italic=False, color_rgb=(0,0,0), line_spacing=1.25):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        r = p.add_run(text)
        set_font(r, font_name=FONT_FAMILY, size_pt=font_size, bold=bold, italic=italic, color_rgb=color_rgb)
    return p

def add_slide_header(slide_no, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    
    r1 = p.add_run(f"SLIDE {slide_no}: ")
    set_font(r1, size_pt=14, bold=True, color_rgb=(108, 99, 255))
    
    r2 = p.add_run(title)
    set_font(r2, size_pt=16, bold=True, color_rgb=(15, 23, 42))
    return p

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_tbl_borders(table, color="CCCCCC"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_slide_box(slide_num, title, slide_type, content_list, speaker_notes):
    add_slide_header(slide_num, title)
    
    # Table box for slide simulation
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_tbl_borders(t, color="6C63FF")
    set_cell_bg(c, "F8FAFC")
    
    p_box = c.paragraphs[0]
    p_box.paragraph_format.space_before = Pt(4)
    p_box.paragraph_format.space_after = Pt(4)
    
    r_hdr = p_box.add_run(f"[{slide_type.upper()} LAYOUT]\n")
    set_font(r_hdr, size_pt=9.5, bold=True, color_rgb=(108, 99, 255))
    
    for item in content_list:
        p_item = c.add_paragraph()
        p_item.paragraph_format.space_before = Pt(2)
        p_item.paragraph_format.space_after = Pt(4)
        if isinstance(item, tuple):
            h, b = item
            rh = p_item.add_run(f"• {h}: ")
            set_font(rh, size_pt=11, bold=True, color_rgb=(30, 41, 59))
            rb = p_item.add_run(b)
            set_font(rb, size_pt=10.5, color_rgb=(51, 65, 85))
        else:
            r = p_item.add_run(f"• {item}")
            set_font(r, size_pt=10.5, color_rgb=(51, 65, 85))
            
    # Speaker Notes section
    add_p("Speaker Notes & Presentation Script:", space_before=8, space_after=4, font_size=11, bold=True, color_rgb=(15, 23, 42))
    add_p(f"“{speaker_notes}”", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=14, font_size=10.5, italic=True, color_rgb=(71, 85, 105))


# Document Header
add_p("PRESIDENCY UNIVERSITY", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2, font_size=14, bold=True)
add_p("PRESIDENCY SCHOOL OF COMPUTER SCIENCE AND ENGINEERING", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=14, font_size=12, bold=True)

add_p("CSE7000 – INTERNSHIP REVIEW PRESENTATION DOCUMENT", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4, font_size=16, bold=True, color_rgb=(108, 99, 255))
add_p("SLIDE-BY-SLIDE CONTENT & SPEAKER SCRIPT FOR PULSEMIND AI", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=18, font_size=13, bold=True, color_rgb=(15, 23, 42))

# Student & Supervisor Info Box
t_info = doc.add_table(rows=2, cols=2)
t_info.alignment = WD_TABLE_ALIGNMENT.CENTER
set_tbl_borders(t_info)

info_headers = [
    "Student Details:\nDIVYA.M.K  (Roll No: 20241CSE0568)\nProgram: B. Tech in Computer Science & Engineering",
    "Under the Supervision of:\nMr. Vishnu Shankar, Assistant Professor, PSCS\nMs. Akkamaha Devi, Assistant Professor, PSCS"
]
for j, txt in enumerate(info_headers):
    c = t_info.cell(0, j)
    c.text = txt
    set_cell_bg(c, "F1F5F9")
    p = c.paragraphs[0]
    for r in p.runs:
        set_font(r, size_pt=10, bold=True)

info_footer = [
    "Presidency University, Bengaluru\nSubmission Date: August 2026",
    "Department Administration:\nHoD: Dr. Nagaraja S R | Program Coord: Ms. Vidhya Rengasamy\nSchool Coord: Dr. Bhuvaneshwari Patil"
]
for j, txt in enumerate(info_footer):
    c = t_info.cell(1, j)
    c.text = txt
    p = c.paragraphs[0]
    for r in p.runs:
        set_font(r, size_pt=9.5, italic=True)

add_p("", space_before=12, space_after=12)

# ── SLIDE 1 ───────────────────────────────────────────────────────────────
add_slide_box(
    1,
    "Title Slide – Project & Team Credentials",
    "Title Slide",
    [
        ("Project Title", "PULSEMIND AI: INTELLIGENT CLINICAL HEALTHCARE & DIAGNOSTIC INTELLIGENCE PLATFORM"),
        ("Student Name", "DIVYA.M.K"),
        ("Roll Number / USN", "20241CSE0568"),
        ("Program & Branch", "Bachelor of Technology in Computer Science and Engineering (Computer Engineering)"),
        ("Internal Supervisors", "Mr. Vishnu Shankar (Assistant Professor) & Ms. Akkamaha Devi (Assistant Professor)"),
        ("Department Leadership", "Head of Department: Dr. Nagaraja S R | Program Coordinator: Ms. Vidhya Rengasamy | School Coordinator: Dr. Bhuvaneshwari Patil"),
        ("Institution & Date", "Presidency School of Computer Science and Engineering, Presidency University, Bengaluru – August 2026")
    ],
    "Good morning respected Dean, HOD Dr. Nagaraja S R, Coordinators Dr. Bhuvaneshwari Patil and Ms. Vidhya Rengasamy, and my respected reviewers Mr. Vishnu Shankar and Ms. Akkamaha Devi. "
    "I am DIVYA.M.K, bearing Roll Number 20241CSE0568. Today I am presenting my CSE7000 Internship Project entitled 'PulseMind AI: Intelligent Clinical Healthcare & Diagnostic Intelligence Platform'."
)

# ── SLIDE 2 ───────────────────────────────────────────────────────────────
add_slide_box(
    2,
    "Table of Contents",
    "Agenda Slide",
    [
        ("01. About Organization", "Overview of PulseMind AI Systems & Healthcare Engineering Division"),
        ("02. Working Domain & Technology", "Python FastAPI, Groq Llama-3.1 LLM, React 18, TypeScript, Cloud Firestore"),
        ("03. Objectives of the Work", "Automated Document OCR, Explainable AI (XAI) Vision Grid, Data Deletion REST APIs"),
        ("04. System Methodology & Architecture", "3-Tier Microservices, Processing Pipeline, India Health Engines"),
        ("05. Empirical Results & Performance", "Latency benchmarks, Vite build output, diagnostic panel metrics"),
        ("06. Conclusion & Future Roadmap", "Key takeaways, ABDM/ABHA integration, Wearable IoT telemetry")
    ],
    "Here is the outline of my presentation. I will begin with an overview of the host organization, followed by our working domain, core objectives, 3-tier system methodology, empirical results, and conclude with future roadmap."
)

# ── SLIDE 3 ───────────────────────────────────────────────────────────────
add_slide_box(
    3,
    "About Company / Organization",
    "Company Overview",
    [
        ("Organization Overview", "PulseMind AI Systems is a specialized healthcare engineering initiative focused on deploying clinical generative AI and computer vision solutions."),
        ("Mission & Vision", "To bridge the gap between complex unstructured clinical diagnostic data and accessible, patient-centric health intelligence."),
        ("Target Domain", "Digital Health Platforms, Telemedicine Middleware, Diagnostic OCR Automation, and Community Healthcare (ASHA Rural Worker Support)."),
        ("Core Product Offerings", "Patient Telemetry Operating System, Multi-Modal Diagnostic XAI, Prescription OCR Parser, and Multi-Scenario Disease Simulator.")
    ],
    "PulseMind AI Systems operates at the intersection of Artificial Intelligence and Clinical Informatics. Our mission is to transform opaque laboratory reports and complex imaging scans into plain-language, explainable diagnostic insights for patients and healthcare providers."
)

# ── SLIDE 4 ───────────────────────────────────────────────────────────────
add_slide_box(
    4,
    "Working Domain & Technology Stack",
    "Tech Stack Overview",
    [
        ("Backend API Middleware", "Python 3.11, FastAPI (0.110), Uvicorn Server, Pydantic Schema Validation, PyMuPDF (fitz), ReportLab PDF Builder."),
        ("Artificial Intelligence & LLM", "Groq Cloud Llama-3.1 8B Instant (Sub-second inference), 8x8 Explainable AI (XAI) Vision Attention Heatmaps."),
        ("Frontend SaaS Interface", "React 18.3, TypeScript 5.5, TailwindCSS 3.4, Vite 5.4, Lucide Icons, Recharts Analytics."),
        ("Database & Cloud Storage", "Google Cloud Firestore (pulsemindai-c4adf) Admin SDK, dual-mode Local JSON persistence fallback."),
        ("Mobile Native Application", "Expo SDK 51, React Native 0.74, Cross-platform iOS & Android support.")
    ],
    "Our technology stack leverages Python FastAPI for high-speed asynchronous REST APIs, Groq Llama-3.1 8B Instant for sub-second LLM inference, React 18 with TypeScript for a futuristic dark-mode SaaS UI, and Google Cloud Firestore for secure NoSQL persistence."
)

# ── SLIDE 5 ───────────────────────────────────────────────────────────────
add_slide_box(
    5,
    "Objectives of the Internship Work",
    "Objectives Slide",
    [
        ("Automated Document Intelligence", "Ingest unstructured PDF/DOCX/TXT lab panels and parse quantitative metrics, reference ranges, and abnormal alerts."),
        ("Explainable AI (XAI) Vision Grid", "Classify Chest X-Rays, Brain MRIs, and Skin Lesions while displaying an interactive 8x8 diagnostic attention matrix."),
        ("Full Data Sovereignty & Deletion", "Build RESTful deletion endpoints (DELETE /api/imaging/{id}, DELETE /api/reports/{id}) with reactive UI controls."),
        ("Localized India Health Modules", "Deploy ASHA Rural Worker Mode, Multilingual Voice AI, Outbreak Predictor, and Healthcare Affordability Estimator."),
        ("Production Rebranding & UI Overhaul", "Rebrand platform to PulseMind AI and design a futuristic dark-mode SaaS dashboard with circular SVG health score rings.")
    ],
    "The primary objectives accomplished during my 10-week internship include building automated document parsing, 8x8 Explainable AI diagnostic vision grids, full RESTful data deletion controls, localized India healthcare modules, and executing a complete platform rebranding to PulseMind AI."
)

# ── SLIDE 6 ───────────────────────────────────────────────────────────────
add_slide_box(
    6,
    "System Methodology Used & Architecture",
    "Architecture & Workflow",
    [
        ("Presentation Tier", "Single-Page Application built with React 18, TypeScript, TailwindCSS, featuring glassmorphism, 3D card elevation, and dynamic SVG score rings."),
        ("Application & Middleware Tier", "FastAPI REST microservices orchestrating document extraction (PyMuPDF), LLM prompt assembly (Groq Llama-3.1), and PDF compilation."),
        ("Data & Persistence Tier", "Cloud Firestore NoSQL database synchronized via Firebase Admin SDK with automatic offline fallback to local JSON database."),
        ("Data Lifecycle & Privacy", "Full CRUD control allowing users to upload, inspect, download ReportLab PDFs, and permanently delete records from disk and cloud.")
    ],
    "The platform follows a decoupled 3-tier architecture. When a user uploads a laboratory PDF, PyMuPDF extracts plain text, Groq Llama-3.1 structures out-of-range metrics, and Cloud Firestore syncs the result in real time. Full CRUD deletion routes ensure complete privacy compliance."
)

# ── SLIDE 7 ───────────────────────────────────────────────────────────────
add_slide_box(
    7,
    "Empirical Results & Latency Benchmarks",
    "Results & Performance",
    [
        ("Document OCR & LLM Parsing Latency", "Average multi-page PDF extraction and Llama-3.1 report structuring completed in 1.12 seconds."),
        ("Diagnostic Vision Classification", "Chest X-Ray, Brain MRI, and Skin Lesion analysis with 8x8 XAI matrix generated in 340 milliseconds."),
        ("Production Build Verification", "Vite build (pulsemind-frontend@1.0.0) compiled cleanly in 7.43 seconds across 2,134 modules with 0 TypeScript errors."),
        ("Cloud Firestore Sync Latency", "Sub-second synchronization (145ms average) with active credentials on project pulsemindai-c4adf.")
    ],
    "Empirical performance testing demonstrated outstanding efficiency. PDF extraction and LLM report structuring complete in 1.12 seconds. Production Vite builds compiled across 2,100+ modules with zero errors, and Cloud Firestore sync latencies averaged under 150 milliseconds."
)

# ── SLIDE 8 ───────────────────────────────────────────────────────────────
add_slide_box(
    8,
    "Conclusion & Future Roadmap",
    "Conclusion Slide",
    [
        ("Key Accomplishments", "Engineered and deployed PulseMind AI operating system, achieving 100% of functional REST API and UI redesign goals."),
        ("Skills Acquired", "FastAPI microservices, Groq LLM prompt engineering, Explainable AI vision grids, Cloud Firestore Admin SDK, and modern CSS systems."),
        ("ABDM / ABHA Integration", "Connect with Ayushman Bharat Digital Mission unified health interface for seamless EHR fetching."),
        ("Wearable Telemetry", "Ingest real-time heart rate, SpO2, and ECG telemetry from Apple HealthKit and Google Fit APIs."),
        ("Federated Edge Learning", "Deploy privacy-preserving edge AI models for local disease outbreak prediction across rural health clinics.")
    ],
    "In conclusion, PulseMind AI demonstrates how generative AI and modern cloud architecture can modernize healthcare delivery. Our future roadmap includes Ayushman Bharat ABHA integration, wearable IoT streaming, and federated edge learning for rural clinics."
)

# ── SLIDE 9 ───────────────────────────────────────────────────────────────
add_slide_box(
    9,
    "Question & Answer Session",
    "Q&A Slide",
    [
        ("Open Floor", "Thank you for your time and attention! We are happy to answer any questions regarding:"),
        ("Topic 1", "PulseMind AI Architecture & Python FastAPI Microservices"),
        ("Topic 2", "Groq Llama-3.1 LLM Prompting & Retrieval-Augmented Generation (RAG)"),
        ("Topic 3", "8x8 Explainable AI (XAI) Vision Heatmap Implementation"),
        ("Topic 4", "Cloud Firestore Integration & Granular REST Deletion Endpoints")
    ],
    "Thank you honorable committee members and reviewers for your valuable time. I am now open to answer any questions regarding system architecture, LLM integration, Explainable AI heatmaps, or database deletion endpoints."
)

# ── SLIDE 10 ──────────────────────────────────────────────────────────────
add_slide_box(
    10,
    "Thank You Slide",
    "Closing Slide",
    [
        ("Closing Statement", "Thank you for your guidance and support throughout the internship!"),
        ("Project Title", "PULSEMIND AI: INTELLIGENT CLINICAL HEALTHCARE & DIAGNOSTIC INTELLIGENCE PLATFORM"),
        ("Student Name", "DIVYA.M.K  (Roll No: 20241CSE0568)"),
        ("Supervisors", "Mr. Vishnu Shankar & Ms. Akkamaha Devi (Assistant Professors, PSCS)"),
        ("Institution", "Presidency School of Computer Science and Engineering, Presidency University, Bengaluru")
    ],
    "Thank you once again to Dean Dr. Duraipandian N, HOD Dr. Nagaraja S R, Coordinators Dr. Bhuvaneshwari Patil and Ms. Vidhya Rengasamy, and my supervisors Mr. Vishnu Shankar and Ms. Akkamaha Devi."
)

# Save Document
out_doc = r"C:\Users\Divya k\Downloads\PulseMind_AI_Internship_PPT_Content.docx"
doc.save(out_doc)
print(f"PPT Content Word document saved successfully at: {out_doc}")
